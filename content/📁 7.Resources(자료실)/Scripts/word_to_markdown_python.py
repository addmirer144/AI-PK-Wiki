#!/usr/bin/env python3
"""
Word to Markdown Converter for Obsidian (Python-only version)
Pandoc 없이도 작동하는 순수 Python 기반 변환기
"""

import os
import sys
import argparse
from pathlib import Path
from datetime import datetime
import re

def check_python_dependencies():
    """Python 의존성 확인 및 설치"""
    required_packages = ['python-docx', 'mammoth']
    missing_packages = []
    
    for package in required_packages:
        try:
            if package == 'python-docx':
                import docx
            elif package == 'mammoth':
                import mammoth
        except ImportError:
            missing_packages.append(package)
    
    if missing_packages:
        print(f"❌ 필요한 패키지가 설치되어 있지 않습니다: {', '.join(missing_packages)}")
        print("설치 명령:")
        for package in missing_packages:
            print(f"  pip3 install {package} --user")
        return False
    else:
        print("✅ 모든 Python 의존성이 확인되었습니다.")
        return True

def convert_with_mammoth(word_file, output_file):
    """mammoth를 사용한 변환 (권장)"""
    try:
        import mammoth
        
        with open(word_file, "rb") as docx_file:
            result = mammoth.convert_to_markdown(docx_file)
            
        # 변환된 마크다운 내용 정리
        markdown_content = clean_markdown_content(result.value)
        
        # 파일 저장
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        # 경고 메시지가 있으면 출력
        if result.messages:
            print(f"⚠️ 변환 경고:")
            for message in result.messages:
                print(f"   {message}")
                
        return True
        
    except ImportError:
        return False
    except Exception as e:
        print(f"❌ Mammoth 변환 실패: {e}")
        return False

def convert_with_docx(word_file, output_file):
    """python-docx를 사용한 기본 변환"""
    try:
        from docx import Document
        
        doc = Document(word_file)
        markdown_lines = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # 스타일에 따른 마크다운 변환
            if paragraph.style.name.startswith('Heading'):
                level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                markdown_lines.append('#' * level + ' ' + text)
            else:
                # 일반 텍스트
                if text:
                    markdown_lines.append(text)
            
            markdown_lines.append('')  # 빈 줄 추가
        
        # 테이블 처리
        for table in doc.tables:
            markdown_lines.append('| ' + ' | '.join([cell.text.strip() for cell in table.rows[0].cells]) + ' |')
            markdown_lines.append('| ' + ' | '.join(['---' for _ in table.rows[0].cells]) + ' |')
            
            for row in table.rows[1:]:
                markdown_lines.append('| ' + ' | '.join([cell.text.strip() for cell in row.cells]) + ' |')
            
            markdown_lines.append('')
        
        markdown_content = '\n'.join(markdown_lines)
        markdown_content = clean_markdown_content(markdown_content)
        
        # 파일 저장
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
            
        return True
        
    except ImportError:
        print("❌ python-docx가 설치되어 있지 않습니다.")
        return False
    except Exception as e:
        print(f"❌ python-docx 변환 실패: {e}")
        return False

def clean_markdown_content(content):
    """마크다운 내용 정리"""
    # 연속된 빈 줄 제거
    content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
    
    # 앞뒤 공백 제거
    content = content.strip()
    
    # 잘못된 마크다운 문법 수정
    content = re.sub(r'^(\s*)(#+)\s*(.+)', r'\1\2 \3', content, flags=re.MULTILINE)
    
    return content

def find_word_files(vault_path):
    """볼트에서 Word 파일 찾기"""
    word_files = []
    vault_path = Path(vault_path)
    
    # .docx와 .doc 파일 검색
    for ext in ['*.docx', '*.doc']:
        word_files.extend(vault_path.rglob(ext))
    
    return sorted(word_files)

def sanitize_filename(filename):
    """파일명을 Obsidian에 적합하게 정리"""
    import re
    name = re.sub(r'[<>:"/\\|?*]', '', filename)
    name = re.sub(r'\s+', ' ', name).strip()
    return name

def convert_word_to_markdown(word_file, output_dir=None):
    """Word 파일을 마크다운으로 변환"""
    word_file = Path(word_file)
    
    # 출력 디렉토리 설정
    if output_dir is None:
        output_dir = word_file.parent / "📁 2.Processed_Notes(가공된노트)"
    else:
        output_dir = Path(output_dir)
    
    # 출력 디렉토리 생성
    output_dir.mkdir(exist_ok=True)
    
    # 출력 파일명 생성
    base_name = sanitize_filename(word_file.stem)
    output_file = output_dir / f"{base_name}.md"
    
    print(f"🔄 변환 중: {word_file.name}")
    
    # 변환 시도 (mammoth 우선, 실패하면 python-docx)
    success = False
    
    # 1순위: mammoth 사용
    success = convert_with_mammoth(word_file, output_file)
    
    # 2순위: python-docx 사용
    if not success:
        print("   mammoth 실패, python-docx로 재시도...")
        success = convert_with_docx(word_file, output_file)
    
    if success:
        # 변환된 파일에 메타데이터 추가
        add_frontmatter(output_file, word_file)
        print(f"✅ 변환 완료: {output_file.name}")
        return output_file
    else:
        print(f"❌ 변환 실패: {word_file.name}")
        return None

def add_frontmatter(markdown_file, original_word_file):
    """마크다운 파일에 Obsidian 메타데이터 추가"""
    markdown_file = Path(markdown_file)
    original_word_file = Path(original_word_file)
    
    # 기존 내용 읽기
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 파일 정보 수집
    created_time = datetime.fromtimestamp(original_word_file.stat().st_birthtime)
    modified_time = datetime.fromtimestamp(original_word_file.stat().st_mtime)
    
    # 태그 추출 (파일명에서)
    filename_lower = original_word_file.name.lower()
    tags = []
    if '새벽' in filename_lower:
        tags.append('#새벽묵상')
    if '청소년' in filename_lower:
        tags.append('#청소년설교')
    if '아침묵상' in filename_lower:
        tags.append('#아침묵상')
    if any(book in filename_lower for book in ['마태복음', '마가복음', '누가복음', '요한복음']):
        tags.append('#복음서')
    if '신명기' in filename_lower or '룻기' in filename_lower:
        tags.append('#구약')
    
    # 기본 태그 추가
    tags.extend(['#설교', '#변환됨'])
    
    # Frontmatter 생성
    frontmatter = f"""---
id: {datetime.now().strftime('%Y%m%d%H%M%S')}
type: processed
created: {created_time.strftime('%Y-%m-%d %H:%M')}
updated: {modified_time.strftime('%Y-%m-%d %H:%M')}
converted: {datetime.now().strftime('%Y-%m-%d %H:%M')}
original_file: {original_word_file.name}
tags: {' '.join(tags)}
---

# {markdown_file.stem}

> **원본**: `{original_word_file.name}`  
> **변환일시**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

"""
    
    # 새로운 내용으로 덮어쓰기
    with open(markdown_file, 'w', encoding='utf-8') as f:
        f.write(frontmatter + content)

def create_conversion_log(converted_files, vault_path):
    """변환 로그 생성"""
    log_file = Path(vault_path) / "word_conversion_log.md"
    
    log_content = f"""# Word to Markdown 변환 로그 (Python 버전)

**변환 일시**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
**변환된 파일 수**: {len(converted_files)}

## 변환된 파일 목록

"""
    
    for i, file_path in enumerate(converted_files, 1):
        log_content += f"{i}. [[{Path(file_path).stem}]]\n"
    
    log_content += f"""

## 다음 단계

- [ ] 변환된 파일들 검토하기
- [ ] 필요한 경우 수동으로 포맷팅 조정
- [ ] 원본 Word 파일 아카이브 또는 삭제 결정
- [ ] 연결 노트 생성하기

## 참고사항

이 변환은 Python 라이브러리를 사용했습니다:
- mammoth: 고품질 HTML/Markdown 변환
- python-docx: 기본적인 텍스트 추출

더 정확한 변환을 원한다면 Pandoc 설치를 고려해보세요:
```bash
brew install pandoc  # macOS
```

---
*자동 생성된 로그 파일*
"""
    
    with open(log_file, 'w', encoding='utf-8') as f:
        f.write(log_content)
    
    print(f"📝 변환 로그 생성: {log_file}")

def main():
    parser = argparse.ArgumentParser(description="Word 파일을 마크다운으로 변환 (Python 버전)")
    parser.add_argument("vault_path", nargs='?', 
                       default="/Users/addmirer/Library/Mobile Documents/iCloud~md~obsidian/Documents/AI-PK M-System",
                       help="Obsidian 볼트 경로")
    parser.add_argument("--file", "-f", help="특정 파일만 변환")
    parser.add_argument("--output", "-o", help="출력 디렉토리 지정")
    parser.add_argument("--dry-run", action="store_true", help="실제 변환 없이 파일 목록만 표시")
    parser.add_argument("--install-deps", action="store_true", help="필요한 Python 패키지 설치")
    
    args = parser.parse_args()
    
    print("🚀 Word to Markdown 변환기 시작 (Python 버전)")
    print("="*50)
    
    # 의존성 설치 모드
    if args.install_deps:
        print("📦 Python 패키지 설치 중...")
        import subprocess
        packages = ['python-docx', 'mammoth']
        for package in packages:
            try:
                subprocess.run([sys.executable, '-m', 'pip', 'install', package, '--user'], check=True)
                print(f"✅ {package} 설치 완료")
            except subprocess.CalledProcessError as e:
                print(f"❌ {package} 설치 실패: {e}")
        return
    
    # 의존성 확인
    if not check_python_dependencies():
        print("\n설치 후 다시 실행하거나 다음 명령으로 자동 설치:")
        print(f"python3 {__file__} --install-deps")
        sys.exit(1)
    
    # 볼트 경로 확인
    vault_path = Path(args.vault_path)
    if not vault_path.exists():
        print(f"❌ 볼트 경로를 찾을 수 없습니다: {vault_path}")
        sys.exit(1)
    
    print(f"📁 볼트 경로: {vault_path}")
    
    # Word 파일 찾기
    if args.file:
        file_path = Path(args.file)
        # 상대 경로인 경우 볼트 경로와 결합
        if not file_path.is_absolute():
            file_path = vault_path / file_path
        word_files = [file_path]
        if not word_files[0].exists():
            print(f"❌ 파일을 찾을 수 없습니다: {word_files[0]}")
            sys.exit(1)
    else:
        word_files = find_word_files(vault_path)
    
    if not word_files:
        print("📄 변환할 Word 파일이 없습니다.")
        return
    
    print(f"📊 발견된 Word 파일: {len(word_files)}개")
    
    # 파일 목록 표시
    for i, file in enumerate(word_files, 1):
        try:
            relative_path = file.relative_to(vault_path)
            print(f"  {i:2d}. {relative_path}")
        except ValueError:
            # 상대 경로 변환 실패 시 파일명만 표시
            print(f"  {i:2d}. {file.name}")
    
    if args.dry_run:
        print("\n🔍 Dry run 모드: 실제 변환을 수행하지 않습니다.")
        return
    
    # 사용자 확인
    if len(word_files) > 1:
        response = input(f"\n{len(word_files)}개 파일을 모두 변환하시겠습니까? (y/N): ")
        if response.lower() != 'y':
            print("변환을 취소합니다.")
            return
    
    print("\n🔄 변환 시작...")
    print("-" * 50)
    
    # 변환 실행
    converted_files = []
    failed_files = []
    
    for word_file in word_files:
        result = convert_word_to_markdown(word_file, args.output)
        if result:
            converted_files.append(result)
        else:
            failed_files.append(word_file)
    
    # 결과 요약
    print("\n" + "="*50)
    print("🎉 변환 완료!")
    print(f"✅ 성공: {len(converted_files)}개")
    print(f"❌ 실패: {len(failed_files)}개")
    
    if failed_files:
        print("\n실패한 파일들:")
        for file in failed_files:
            print(f"  - {file.name}")
    
    # 변환 로그 생성
    if converted_files:
        create_conversion_log(converted_files, vault_path)
    
    print(f"\n📁 변환된 파일들은 다음 위치에 저장되었습니다:")
    if args.output:
        print(f"  {args.output}")
    else:
        print(f"  각 원본 파일 디렉토리의 '📁 2.Processed_Notes(가공된노트)' 폴더")

if __name__ == "__main__":
    main()